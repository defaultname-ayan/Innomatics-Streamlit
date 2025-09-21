# resume_parser.py - ENHANCED VERSION
import fitz  # PyMuPDF
import docx
import re
from typing import Dict, List, Optional, Tuple
import warnings
import logging

warnings.filterwarnings("ignore")
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ResumeParser:
    def __init__(self):
        # Comprehensive skill database for better accuracy
        self.skill_keywords = {
            # Programming Languages
            "languages": [
                "python", "java", "javascript", "typescript", "c", "c++", "c#", "go", 
                "rust", "scala", "kotlin", "sql", "r", "swift", "dart", "php", "ruby",
                "matlab", "julia", "perl", "shell", "bash", "powershell"
            ],
            # Web Technologies
            "web": [
                "react", "angular", "vue", "svelte", "node.js", "nodejs", "express", 
                "django", "flask", "fastapi", "spring", "next.js", "nuxt", "gatsby",
                "html", "css", "sass", "scss", "bootstrap", "tailwind"
            ],
            # Data Science & AI
            "data_science": [
                "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "keras",
                "xgboost", "lightgbm", "catboost", "mlflow", "dvc", "airflow",
                "jupyter", "matplotlib", "seaborn", "plotly", "tableau", "power bi"
            ],
            # Cloud & DevOps
            "cloud": [
                "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible",
                "jenkins", "github actions", "gitlab ci", "circleci", "helm", "istio",
                "prometheus", "grafana", "elk stack", "datadog"
            ],
            # Databases
            "databases": [
                "mysql", "postgresql", "mongodb", "redis", "cassandra", "dynamodb",
                "snowflake", "bigquery", "elasticsearch", "neo4j", "sqlite", "oracle"
            ],
            # Specialized
            "specialized": [
                "restful api", "graphql", "grpc", "microservices", "blockchain", "nlp",
                "computer vision", "deep learning", "machine learning", "data mining",
                "etl", "ci/cd", "agile", "scrum", "kanban", "tdd", "bdd"
            ]
        }
        
        # Flatten all skills for quick lookup
        self.all_skills = []
        for category in self.skill_keywords.values():
            self.all_skills.extend(category)
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Enhanced PDF text extraction with error handling"""
        try:
            doc = fitz.open(pdf_path)
            text = ""
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if not page_text.strip():
                    # Try OCR if no text found
                    try:
                        pix = page.get_pixmap()
                        # Note: OCR would require pytesseract
                        logger.warning(f"Page {page_num} appears to be image-based")
                    except Exception:
                        pass
                text += page_text + "\n"
            doc.close()
            return text
        except Exception as e:
            logger.error(f"Error reading PDF {pdf_path}: {str(e)}")
            return f"Error reading PDF: {str(e)}"
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """Enhanced DOCX extraction with table support"""
        try:
            doc = docx.Document(docx_path)
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error reading DOCX {docx_path}: {str(e)}")
            return f"Error reading DOCX: {str(e)}"
    
    def extract_skills(self, text: str) -> List[str]:
        """Enhanced skill extraction with normalization and context"""
        text_lower = text.lower()
        found_skills = set()
        
        # Direct keyword matching with word boundaries
        for skill in self.all_skills:
            # Handle multi-word skills
            if len(skill.split()) > 1:
                if skill in text_lower:
                    found_skills.add(skill)
            else:
                # Use word boundaries for single words
                pattern = rf'\b{re.escape(skill)}\b'
                if re.search(pattern, text_lower):
                    found_skills.add(skill)
        
        # Additional pattern matching for variations
        skill_variations = {
            'javascript': ['js', 'node.js', 'nodejs'],
            'python': ['py'],
            'kubernetes': ['k8s'],
            'machine learning': ['ml'],
            'artificial intelligence': ['ai'],
            'natural language processing': ['nlp']
        }
        
        for main_skill, variations in skill_variations.items():
            for var in variations:
                if re.search(rf'\b{re.escape(var)}\b', text_lower):
                    found_skills.add(main_skill)
        
        return sorted(list(found_skills))
    
    def extract_experience(self, text: str) -> str:
        """Enhanced experience extraction"""
        patterns = [
            r'(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)',
            r'(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)',
            r'experience[:\s]+(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text.lower())
            if match:
                return f"{match.group(1)} years"
        
        return "Not specified"
    
    def extract_education(self, text: str) -> str:
        """Enhanced education extraction"""
        text_lower = text.lower()
        
        education_levels = [
            (["phd", "ph.d", "doctorate", "doctoral"], "PhD"),
            (["master", "m.tech", "mtech", "m.sc", "msc", "mba", "m.s"], "Master's"),
            (["bachelor", "b.tech", "btech", "b.e", "be", "b.sc", "bsc", "b.s"], "Bachelor's"),
            (["diploma", "associate"], "Diploma/Associate")
        ]
        
        for keywords, level in education_levels:
            if any(keyword in text_lower for keyword in keywords):
                return level
        
        return "Not specified"
    
    def extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information"""
        contact = {}
        
        # Email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        contact['email'] = email_match.group(0) if email_match else ""
        
        # Phone
        phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phone_match = re.search(phone_pattern, text)
        contact['phone'] = phone_match.group(0) if phone_match else ""
        
        return contact
    
    def parse_resume(self, file_path: str) -> Dict:
        """Main parsing method with comprehensive extraction"""
        try:
            # Extract text based on file type
            if file_path.lower().endswith('.pdf'):
                raw_text = self.extract_text_from_pdf(file_path)
            elif file_path.lower().endswith('.docx'):
                raw_text = self.extract_text_from_docx(file_path)
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    raw_text = f.read()
            
            # Extract structured information
            skills = self.extract_skills(raw_text)
            experience = self.extract_experience(raw_text)
            education = self.extract_education(raw_text)
            contact = self.extract_contact_info(raw_text)
            
            return {
                'raw_text': raw_text,
                'skills': skills,
                'experience': experience,
                'education': education,
                'contact': contact,
                'word_count': len(raw_text.split()),
                'skill_count': len(skills)
            }
            
        except Exception as e:
            logger.error(f"Error parsing resume {file_path}: {str(e)}")
            return {
                'raw_text': f"Error parsing resume: {str(e)}",
                'skills': [],
                'experience': "Not specified",
                'education': "Not specified",
                'contact': {},
                'word_count': 0,
                'skill_count': 0
            }
